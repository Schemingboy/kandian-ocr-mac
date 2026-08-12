#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
看典古籍 OCR · Mac 客户端 v1.0
================================================
复刻 Windows 客户端「看典古籍OCR客户端v2.0.7」的识别流程，供 macOS 使用。

  识别流程（与 Windows 版一致）：
    PDF 识别    本地逐页把 PDF 转成图片，逐页调用云端 OCR 接口，逐页保存、断点可续、末尾汇总
    单图识别    选一张古籍图片直接识别
    额度查询    查询 API Token 剩余额度与状态

  云端接口：https://ocr.kandianguji.com/ocr_api  （token + email 认证）

  首次使用：先运行 install.sh 安装依赖，之后双击 run.command 启动。
  依赖：PySide6、requests、PyMuPDF
"""

import base64
import json
import os
import subprocess
import sys
from pathlib import Path

from PySide6 import QtCore, QtWidgets

API_URL = "https://ocr.kandianguji.com/ocr_api"
STATUS_URL = "https://ocr.kandianguji.com/get_token_status"
CONFIG_FILE = Path.home() / ".kandian_ocr.json"
DEFAULT_EMAIL = ""  # 注册账号的邮箱/手机号，由用户在界面里填写，不写死在代码里
IS_MAC = sys.platform == "darwin"


# ============================ 配置 ============================
def load_config():
    cfg = {
        "email": DEFAULT_EMAIL, "token": "",
        "version": "v2", "det_mode": "auto", "image_size": 1024,
        "only_plain_text": True, "resume": True,
    }
    try:
        if CONFIG_FILE.exists():
            cfg.update(json.loads(CONFIG_FILE.read_text(encoding="utf-8")))
    except Exception:
        pass
    return cfg


def save_config(cfg):
    CONFIG_FILE.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


# ===================== 核心 API（不依赖界面，可单独测试） =====================
def _post_form(url, fields, timeout=120):
    """POST 表单，返回解析后的 dict。网络/解析失败也返回 error dict，绝不抛异常。"""
    try:
        import requests
    except ImportError:
        return {"message": "error", "info": "缺少 requests 依赖，请先运行 install.sh"}
    try:
        r = requests.post(url, data=fields, timeout=timeout)
    except Exception as e:
        return {"message": "error", "info": f"网络请求失败：{e}"}
    try:
        return r.json()
    except Exception:
        return {"message": "error", "info": f"返回异常（HTTP {r.status_code}）：{r.text[:200]}"}


def ocr_image_bytes(token, email, img_bytes, version="v2", det_mode="auto",
                    image_size=1024, only_plain_text=True):
    """单张图片 OCR。返回 (ok: bool, text_or_err: str)。"""
    b64 = base64.b64encode(img_bytes).decode("ascii")
    fields = {
        "token": token, "email": email, "image": b64,
        "version": version, "det_mode": det_mode, "image_size": str(image_size),
    }
    if only_plain_text:
        fields["only_plain_text"] = "true"
    j = _post_form(API_URL, fields)
    if j.get("message") == "success":
        data = j.get("data")
        lines = []
        if isinstance(data, list):
            # 实测真实格式：data 直接就是文本行数组
            lines = [t for t in data if isinstance(t, str)]
        elif isinstance(data, dict):
            # 兼容官方文档描述的 data.texts 格式
            texts = data.get("texts") or []
            lines = [t for t in texts if isinstance(t, str)]
        return True, "\n".join(lines)
    return False, str(j.get("info") or j.get("message") or "未知错误")


def check_quota(token, email):
    """查询 token 额度。返回 (ok: bool, 说明文本)。"""
    j = _post_form(STATUS_URL, {"token": token, "email": email}, timeout=30)
    if j.get("message") != "success":
        return False, str(j.get("info") or j.get("message") or "查询失败")
    d = j.get("data") or {}
    state = {0: "申请中", 1: "已通过 · 可正常使用", 2: "申请未通过"}.get(
        d.get("is_active"), str(d.get("is_active", "?")))
    return True, f"Token 状态：{state} ｜ 已用 {d.get('used_count', '?')} / 总额 {d.get('total_count', '?')}"


# ============================ 汇总文件 ============================
def write_summary_txt(page_dir):
    """把 page_*.txt 合并成 汇总.txt。成功返回 True，失败返回 False。"""
    try:
        parts = []
        for p in sorted(Path(page_dir).glob("page_*.txt")):
            page_no = int(p.stem.split("_")[1])
            parts.append(f"【第 {page_no} 页】")
            parts.append(p.read_text(encoding="utf-8").rstrip("\n"))
            parts.append("")
        (Path(page_dir) / "汇总.txt").write_text("\n".join(parts), encoding="utf-8")
        return True
    except Exception:
        return False


def write_summary_docx(page_dir):
    """把 page_*.txt 合并成 汇总.docx（Word）。python-docx 未装或出错时静默返回 False。"""
    try:
        import docx
    except ImportError:
        return False
    try:
        d = docx.Document()
        d.add_heading("汇总", level=1)
        for p in sorted(Path(page_dir).glob("page_*.txt")):
            page_no = int(p.stem.split("_")[1])
            d.add_heading(f"第 {page_no} 页", level=2)
            for ln in p.read_text(encoding="utf-8").splitlines():
                if ln.strip():
                    d.add_paragraph(ln)
        d.save(str(Path(page_dir) / "汇总.docx"))
        return True
    except Exception:
        return False


# ============================ PDF 识别线程 ============================
class PdfOcrThread(QtCore.QThread):
    sig_progress = QtCore.Signal(int, int, str)          # 当前页, 总页数, 状态文字
    sig_log = QtCore.Signal(str)
    sig_done = QtCore.Signal(int, int, str)              # 成功页数, 失败页数, 输出目录
    sig_fatal = QtCore.Signal(str)

    def __init__(self, token, email, pdf_path, out_dir, params, parent=None):
        super().__init__(parent)
        self.token = token
        self.email = email
        self.pdf_path = pdf_path
        self.out_dir = Path(out_dir)
        self.params = params
        self._stop = False

    def stop(self):
        self._stop = True

    def run(self):
        try:
            import fitz  # PyMuPDF
        except ImportError:
            self.sig_fatal.emit("缺少 PyMuPDF 依赖，请先运行 install.sh")
            return

        try:
            pdf = fitz.open(self.pdf_path)
        except Exception as e:
            self.sig_fatal.emit(f"无法打开 PDF：{e}")
            return

        pdf_name = Path(self.pdf_path).stem
        page_dir = self.out_dir / pdf_name
        try:
            page_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            self.sig_fatal.emit(f"无法创建输出目录：{e}")
            pdf.close()
            return

        total = pdf.page_count
        self.sig_log.emit(f"开始识别：{pdf_name}.pdf，共 {total} 页")
        self.sig_log.emit(f"输出目录：{page_dir}")

        ok_count = skip_count = fail_count = 0
        fail_list = []
        resume = bool(self.params.get("resume", True))

        for i in range(total):
            if self._stop:
                self.sig_log.emit("已手动停止。")
                break
            page_no = i + 1
            page_file = page_dir / f"page_{page_no:04d}.txt"

            if resume and page_file.exists():
                skip_count += 1
                self.sig_progress.emit(page_no, total, f"第 {page_no}/{total} 页：已有结果，跳过")
                continue

            try:
                pix = pdf.load_page(i).get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
            except Exception as e:
                fail_count += 1
                fail_list.append(f"第 {page_no} 页：转图失败 {e}")
                self.sig_log.emit(f"第 {page_no} 页转图失败：{e}")
                self.sig_progress.emit(page_no, total, f"第 {page_no}/{total} 页：转图失败")
                continue

            ok, text = ocr_image_bytes(
                self.token, self.email, img_bytes,
                version=self.params.get("version", "v2"),
                det_mode=self.params.get("det_mode", "auto"),
                image_size=int(self.params.get("image_size", 1024)),
                only_plain_text=bool(self.params.get("only_plain_text", True)),
            )
            if ok:
                try:
                    page_file.write_text(text + "\n", encoding="utf-8")
                except Exception as e:
                    fail_count += 1
                    fail_list.append(f"第 {page_no} 页：保存失败 {e}")
                    self.sig_log.emit(f"第 {page_no} 页保存失败：{e}")
                    continue
                ok_count += 1
                preview = text.strip().replace("\n", " ")[:40]
                self.sig_log.emit(f"第 {page_no} 页 OK：{preview}…")
            else:
                fail_count += 1
                fail_list.append(f"第 {page_no} 页：{text}")
                self.sig_log.emit(f"第 {page_no} 页识别失败：{text}")
            self.sig_progress.emit(page_no, total, f"第 {page_no}/{total} 页")

        pdf.close()

        summary_txt = page_dir / "汇总.txt"
        if write_summary_txt(page_dir):
            self.sig_log.emit(f"已生成文本汇总：{summary_txt.name}")
        if write_summary_docx(page_dir):
            self.sig_log.emit("已生成 Word 汇总：汇总.docx")
        else:
            self.sig_log.emit("（未生成 Word：python-docx 未安装或出错；txt 汇总不受影响）")

        self.sig_log.emit(
            f"完成：成功 {ok_count} 页，跳过 {skip_count} 页，失败 {fail_count} 页。结果目录：{page_dir}")
        for f in fail_list:
            self.sig_log.emit("  ✗ " + f)
        self.sig_done.emit(ok_count, fail_count, str(page_dir))


class ImageOcrThread(QtCore.QThread):
    sig = QtCore.Signal(bool, str)                      # ok, text_or_err

    def __init__(self, token, email, img_bytes, params, parent=None):
        super().__init__(parent)
        self.token, self.email, self.img_bytes, self.params = token, email, img_bytes, params

    def run(self):
        ok, text = ocr_image_bytes(
            self.token, self.email, self.img_bytes,
            version=self.params.get("version", "v2"),
            det_mode=self.params.get("det_mode", "auto"),
            image_size=int(self.params.get("image_size", 1024)),
            only_plain_text=bool(self.params.get("only_plain_text", True)),
        )
        self.sig.emit(ok, text)


class QuotaThread(QtCore.QThread):
    sig = QtCore.Signal(bool, str)

    def __init__(self, token, email, parent=None):
        super().__init__(parent)
        self.token, self.email = token, email

    def run(self):
        ok, msg = check_quota(self.token, self.email)
        self.sig.emit(ok, msg)


# ============================ 界面 ============================
class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("看典古籍 OCR · Mac")
        self.cfg = load_config()
        self.thread = None
        self.pdf_path = ""
        self.out_dir = ""
        self.img_path = ""
        self._last_out = ""

        tabs = QtWidgets.QTabWidget()
        self.setCentralWidget(tabs)
        tabs.addTab(self._build_settings_tab(), "设置")
        tabs.addTab(self._build_pdf_tab(), "PDF 识别")
        tabs.addTab(self._build_image_tab(), "单图识别")
        self.resize(700, 640)
        self._load_cfg_to_ui()

    # ---------- 设置页 ----------
    def _build_settings_tab(self):
        w = QtWidgets.QWidget()
        form = QtWidgets.QFormLayout(w)

        self.ed_email = QtWidgets.QLineEdit()
        self.ed_token = QtWidgets.QLineEdit()
        self.ed_token.setEchoMode(QtWidgets.QLineEdit.Password)
        form.addRow("邮箱 / 手机号：", self.ed_email)
        form.addRow("API Token：", self.ed_token)

        self.ck_show_token = QtWidgets.QCheckBox("显示 Token")
        self.ck_show_token.toggled.connect(
            lambda on: self.ed_token.setEchoMode(
                QtWidgets.QLineEdit.Normal if on else QtWidgets.QLineEdit.Password))
        form.addRow("", self.ck_show_token)

        row = QtWidgets.QHBoxLayout()
        self.btn_save = QtWidgets.QPushButton("保存设置")
        self.btn_save.clicked.connect(self.save_settings)
        self.btn_quota = QtWidgets.QPushButton("检查额度")
        self.btn_quota.clicked.connect(self.start_quota_check)
        row.addWidget(self.btn_save)
        row.addWidget(self.btn_quota)
        form.addRow(row)

        self.lbl_quota = QtWidgets.QLabel("")
        self.lbl_quota.setWordWrap(True)
        form.addRow(self.lbl_quota)

        tip = QtWidgets.QLabel(
            "Token 在官网「古籍数字化 → OCR API」页面获取；邮箱/手机号填你注册账号的那个。\n"
            "设置保存在 ~/.kandian_ocr.json（本机，不联网上传）。")
        tip.setWordWrap(True)
        tip.setStyleSheet("color:#888; font-size:12px;")
        form.addRow(tip)
        return w

    # ---------- PDF 页 ----------
    def _build_pdf_tab(self):
        w = QtWidgets.QWidget()
        v = QtWidgets.QVBoxLayout(w)

        row1 = QtWidgets.QHBoxLayout()
        self.btn_pdf = QtWidgets.QPushButton("选择 PDF 文件…")
        self.btn_pdf.clicked.connect(self.pick_pdf)
        self.lbl_pdf = QtWidgets.QLabel("（未选择）")
        self.lbl_pdf.setWordWrap(True)
        row1.addWidget(self.btn_pdf)
        row1.addWidget(self.lbl_pdf, 1)
        v.addLayout(row1)

        row2 = QtWidgets.QHBoxLayout()
        self.btn_out = QtWidgets.QPushButton("选择保存文件夹…")
        self.btn_out.clicked.connect(self.pick_out)
        self.lbl_out = QtWidgets.QLabel("（未选择）")
        self.lbl_out.setWordWrap(True)
        row2.addWidget(self.btn_out)
        row2.addWidget(self.lbl_out, 1)
        v.addLayout(row2)

        p = QtWidgets.QHBoxLayout()
        p.addWidget(QtWidgets.QLabel("版本："))
        self.cb_version = QtWidgets.QComboBox()
        self.cb_version.addItems(["v2", "default", "beta"])
        p.addWidget(self.cb_version)
        p.addSpacing(8)
        p.addWidget(QtWidgets.QLabel("排版："))
        self.cb_det = QtWidgets.QComboBox()
        self.cb_det.addItem("自动识别", "auto")
        self.cb_det.addItem("横排", "hp")
        self.cb_det.addItem("竖排", "sp")
        p.addWidget(self.cb_det)
        p.addSpacing(8)
        p.addWidget(QtWidgets.QLabel("缩放："))
        self.sp_size = QtWidgets.QSpinBox()
        self.sp_size.setRange(0, 2000)
        self.sp_size.setToolTip("识别前按最长边等比缩放，0 为不调整，建议 1000-2000")
        p.addWidget(self.sp_size)
        p.addStretch(1)
        v.addLayout(p)

        c = QtWidgets.QHBoxLayout()
        self.ck_plain = QtWidgets.QCheckBox("只识别正文")
        self.ck_resume = QtWidgets.QCheckBox("跳过已有结果（断点续跑）")
        c.addWidget(self.ck_plain)
        c.addWidget(self.ck_resume)
        c.addStretch(1)
        v.addLayout(c)

        op = QtWidgets.QHBoxLayout()
        self.btn_start = QtWidgets.QPushButton("开始识别")
        self.btn_start.clicked.connect(self.start_pdf)
        self.btn_stop = QtWidgets.QPushButton("停止")
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self.stop_pdf)
        op.addWidget(self.btn_start)
        op.addWidget(self.btn_stop)
        op.addStretch(1)
        v.addLayout(op)

        self.lbl_prog = QtWidgets.QLabel("")
        v.addWidget(self.lbl_prog)
        self.progress = QtWidgets.QProgressBar()
        self.progress.setFormat("%v / %m 页")
        v.addWidget(self.progress)

        self.btn_open = QtWidgets.QPushButton("打开结果文件夹")
        self.btn_open.setEnabled(False)
        self.btn_open.clicked.connect(self.open_out)
        v.addWidget(self.btn_open)

        self.log = QtWidgets.QPlainTextEdit()
        self.log.setReadOnly(True)
        v.addWidget(self.log, 1)
        return w

    # ---------- 单图页 ----------
    def _build_image_tab(self):
        w = QtWidgets.QWidget()
        v = QtWidgets.QVBoxLayout(w)

        row = QtWidgets.QHBoxLayout()
        self.btn_img = QtWidgets.QPushButton("选择图片…")
        self.btn_img.clicked.connect(self.pick_image)
        self.lbl_img = QtWidgets.QLabel("（未选择）")
        self.lbl_img.setWordWrap(True)
        row.addWidget(self.btn_img)
        row.addWidget(self.lbl_img, 1)
        v.addLayout(row)

        self.btn_img_ocr = QtWidgets.QPushButton("开始识别")
        self.btn_img_ocr.clicked.connect(self.ocr_image)
        v.addWidget(self.btn_img_ocr)

        self.ed_img_result = QtWidgets.QPlainTextEdit()
        v.addWidget(self.ed_img_result, 1)

        self.btn_copy = QtWidgets.QPushButton("复制结果")
        self.btn_copy.clicked.connect(self.copy_result)
        v.addWidget(self.btn_copy)
        return w

    # ---------- 设置相关 ----------
    def _load_cfg_to_ui(self):
        self.ed_email.setText(self.cfg.get("email", DEFAULT_EMAIL))
        self.ed_token.setText(self.cfg.get("token", ""))
        self.cb_version.setCurrentText(self.cfg.get("version", "v2"))
        idx = self.cb_det.findData(self.cfg.get("det_mode", "auto"))
        self.cb_det.setCurrentIndex(max(0, idx))
        self.sp_size.setValue(int(self.cfg.get("image_size", 1024)))
        self.ck_plain.setChecked(bool(self.cfg.get("only_plain_text", True)))
        self.ck_resume.setChecked(bool(self.cfg.get("resume", True)))

    def _collect_cfg(self):
        self.cfg["email"] = self.ed_email.text().strip()
        self.cfg["token"] = self.ed_token.text().strip()
        self.cfg["version"] = self.cb_version.currentText()
        self.cfg["det_mode"] = self.cb_det.currentData()
        self.cfg["image_size"] = self.sp_size.value()
        self.cfg["only_plain_text"] = self.ck_plain.isChecked()
        self.cfg["resume"] = self.ck_resume.isChecked()

    def save_settings(self):
        self._collect_cfg()
        save_config(self.cfg)
        self.lbl_quota.setText("设置已保存 ✓")

    def start_quota_check(self):
        self._collect_cfg()
        token, email = self.cfg.get("token", ""), self.cfg.get("email", "")
        if not token:
            self.lbl_quota.setText("请先填写 API Token。")
            return
        self.btn_quota.setEnabled(False)
        self.lbl_quota.setText("查询中…")
        self.qt = QuotaThread(token, email)
        self.qt.sig.connect(self._on_quota)
        self.qt.start()

    def _on_quota(self, ok, msg):
        self.btn_quota.setEnabled(True)
        self.lbl_quota.setText(msg)

    # ---------- PDF 相关 ----------
    def pick_pdf(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "选择 PDF", "", "PDF 文件 (*.pdf)")
        if path:
            self.pdf_path = path
            self.lbl_pdf.setText(path)

    def pick_out(self):
        d = QtWidgets.QFileDialog.getExistingDirectory(self, "选择保存文件夹")
        if d:
            self.out_dir = d
            self.lbl_out.setText(d)

    def start_pdf(self):
        self._collect_cfg()
        if not self.pdf_path:
            self.log.appendPlainText("请先选择 PDF 文件。")
            return
        if not self.out_dir:
            self.log.appendPlainText("请先选择保存文件夹。")
            return
        if not self.cfg.get("token"):
            self.log.appendPlainText("请先在「设置」页填写 API Token。")
            return

        params = {
            "version": self.cfg["version"], "det_mode": self.cfg["det_mode"],
            "image_size": self.cfg["image_size"],
            "only_plain_text": self.cfg["only_plain_text"],
            "resume": self.cfg["resume"],
        }
        self.thread = PdfOcrThread(self.cfg["token"], self.cfg["email"],
                                   self.pdf_path, self.out_dir, params)
        self.thread.sig_progress.connect(self._on_progress)
        self.thread.sig_log.connect(self.log.appendPlainText)
        self.thread.sig_done.connect(self._on_pdf_done)
        self.thread.sig_fatal.connect(self._on_pdf_fatal)
        self._set_pdf_busy(True)
        self.thread.start()
        save_config(self.cfg)

    def stop_pdf(self):
        if self.thread:
            self.thread.stop()
            self.btn_stop.setEnabled(False)

    def _set_pdf_busy(self, busy):
        self.btn_start.setEnabled(not busy)
        self.btn_stop.setEnabled(busy)
        self.btn_pdf.setEnabled(not busy)
        self.btn_out.setEnabled(not busy)

    def _on_progress(self, done, total, text):
        self.progress.setRange(0, total)
        self.progress.setValue(done)
        self.lbl_prog.setText(text)

    def _on_pdf_done(self, ok, fail, out_dir):
        self._set_pdf_busy(False)
        self._last_out = out_dir
        self.btn_open.setEnabled(True)
        self.progress.setValue(self.progress.maximum())
        self.lbl_prog.setText(f"完成：成功 {ok} 页，失败 {fail} 页")

    def _on_pdf_fatal(self, msg):
        self._set_pdf_busy(False)
        self.log.appendPlainText("✗ " + msg)

    def open_out(self):
        d = self._last_out or self.out_dir
        if not d:
            return
        try:
            if IS_MAC:
                subprocess.Popen(["open", d])
            else:
                os.startfile(d)  # type: ignore[attr-defined]
        except Exception as e:
            self.log.appendPlainText(f"打开文件夹失败：{e}")

    # ---------- 单图相关 ----------
    def pick_image(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self, "选择图片", "", "图片 (*.png *.jpg *.jpeg *.bmp *.webp *.tif *.tiff)")
        if path:
            self.img_path = path
            self.lbl_img.setText(path)

    def ocr_image(self):
        self._collect_cfg()
        if not self.img_path:
            self.ed_img_result.setPlainText("请先选择图片。")
            return
        if not self.cfg.get("token"):
            self.ed_img_result.setPlainText("请先在「设置」页填写 API Token。")
            return
        try:
            img_bytes = Path(self.img_path).read_bytes()
        except Exception as e:
            self.ed_img_result.setPlainText(f"读取图片失败：{e}")
            return
        self.btn_img_ocr.setEnabled(False)
        self.btn_img_ocr.setText("识别中…")
        params = {
            "version": self.cfg["version"], "det_mode": self.cfg["det_mode"],
            "image_size": self.cfg["image_size"],
            "only_plain_text": self.cfg["only_plain_text"],
        }
        self.img_thread = ImageOcrThread(self.cfg["token"], self.cfg["email"], img_bytes, params)
        self.img_thread.sig.connect(self._on_img_ocr)
        self.img_thread.start()

    def _on_img_ocr(self, ok, text):
        self.btn_img_ocr.setEnabled(True)
        self.btn_img_ocr.setText("开始识别")
        if ok:
            self.ed_img_result.setPlainText(text)
        else:
            self.ed_img_result.setPlainText("识别失败：" + text)

    def copy_result(self):
        QtWidgets.QApplication.clipboard().setText(self.ed_img_result.toPlainText())

    # ---------- 退出 ----------
    def closeEvent(self, event):
        if self.thread is not None and self.thread.isRunning():
            self.thread.stop()
            self.thread.wait(3000)
        self._collect_cfg()
        save_config(self.cfg)
        event.accept()


def main():
    app = QtWidgets.QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
